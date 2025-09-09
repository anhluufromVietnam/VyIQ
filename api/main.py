# --- FastAPI app with SQLAlchemy and auto-thumbnail generation ---

from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from sqlalchemy import create_engine, Column, Integer, String, Float
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from typing import Optional, List
from pydantic import BaseModel
import os
import shutil
from moviepy import VideoFileClip
from PIL import Image
import PyPDF2
import pandas as pds
import torch
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader, Settings
from llama_index.llms.openai import OpenAI
from llama_index.embeddings.openai import OpenAIEmbedding
from openai import OpenAI as OpenAIClient
from llama_index.llms.huggingface import HuggingFaceLLM
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from transformers import AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig
import uuid
import re

app = FastAPI()

from fastapi.staticfiles import StaticFiles
from starlette.responses import FileResponse
from starlette.staticfiles import StaticFiles

class NoCacheStaticFiles(StaticFiles):
    async def get_response(self, path, scope):
        response = await super().get_response(path, scope)
        if isinstance(response, FileResponse):
            response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
        return response

app.mount("/project_files", NoCacheStaticFiles(directory="project_files"), name="project_files")
# Serve everything under project_files as static
#app.mount("/project_files", StaticFiles(directory="project_files"), name="project_files")

origins = [
    "http://127.0.0.1:3000",
    "http://localhost:3000",
    "https://babe-belongs-sites-occasional.trycloudflare.com",
    # hoặc "*" cho mọi domain (ít an toàn)
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Database setup ---
DATABASE_URL = "sqlite:///./projects.db"
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# --- Models ---
class Project(Base):
    __tablename__ = "projects"

    id = Column(Integer, primary_key=True, index=True)
    name = Column(String)
    tag = Column(String)
    description = Column(String)
    video_path = Column(String, nullable=True)
    document_path = Column(String, nullable=True)
    accuracy = Column(Float, default=0.0)
    status = Column(String, default="draft")
    thumbnail = Column(String, nullable=True)

Base.metadata.create_all(bind=engine)

# --- Schemas ---
class ProjectSchema(BaseModel):
    id: int
    name: str
    tag: str
    description: str
    video_path: Optional[str] = None
    document_path: Optional[str] = None
    accuracy: float
    status: str
    thumbnail: Optional[str] = None

    class Config:
        orm_mode = True

# --- Helper ---
def generate_thumbnail(video_path: str, thumb_path: str):
    try:
        clip = VideoFileClip(video_path)
        frame = clip.get_frame(1.0)
        image = Image.fromarray(frame)
        image.save(thumb_path)
        return thumb_path
    except Exception as e:
        print(f"Thumbnail generation failed: {e}")
        return None

# --- Routes ---
@app.get("/projects", response_model=List[ProjectSchema])
def get_projects():
    db = SessionLocal()
    projects = db.query(Project).all()
    return projects
    
from fastapi import HTTPException

@app.post("/projects/{project_id}/upload_video")
def upload_video(project_id: int, file: UploadFile = File(...)):
    db = SessionLocal()
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    project_dir = f"project_files/{project.id}/video"
    os.makedirs(project_dir, exist_ok=True)
    video_path = f"{project_dir}/{file.filename}"

    with open(video_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    project.video_path = video_path

    # thumbnail auto tạo
    try:
        from moviepy import VideoFileClip
        thumbnail_path = f"project_files/{project.id}/thumb.jpg"
        clip = VideoFileClip(video_path)
        clip.save_frame(thumbnail_path, t=1.0)
        project.thumbnail = thumbnail_path
    except Exception as e:
        print("Thumbnail generation failed:", e)

    db.commit()
    return {"message": "Video uploaded", "path": video_path}
    
@app.post("/projects/{project_id}/upload_doc")
def upload_doc(project_id: int, file: UploadFile = File(...)):
    db = SessionLocal()
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    docs_dir = f"project_files/{project.id}/docs"
    os.makedirs(docs_dir, exist_ok=True)
    doc_path = f"{docs_dir}/{file.filename}"

    with open(doc_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    project.document_path = doc_path
    db.commit()
    return {"message": "Document uploaded", "path": doc_path}


@app.post("/projects")
def create_project(
    name: str = Form(...),
    tag: str = Form(...),
    description: str = Form(...),
    status: str = Form("draft"),
    accuracy: float = Form(0.0),
    video: Optional[UploadFile] = File(None),
    document: Optional[UploadFile] = File(None),
):
    db = SessionLocal()
    new_project = Project(name=name, tag=tag, description=description, status=status, accuracy=accuracy)
    db.add(new_project)
    db.commit()
    db.refresh(new_project)

    project_dir = f"project_files/{new_project.id}"
    os.makedirs(f"{project_dir}/video", exist_ok=True)
    os.makedirs(f"{project_dir}/docs", exist_ok=True)
    os.makedirs(f"{project_dir}/log", exist_ok=True)

    if video:
        video_path = f"{project_dir}/video/{video.filename}"
        with open(video_path, "wb") as f:
            shutil.copyfileobj(video.file, f)
        new_project.video_path = video_path

        # Generate thumbnail
        thumb_path = f"{project_dir}/video/thumbnail.jpg"
        thumb = generate_thumbnail(video_path, thumb_path)
        if thumb:
            new_project.thumbnail = thumb

    if document:
        doc_path = f"{project_dir}/docs/{document.filename}"
        with open(doc_path, "wb") as f:
            shutil.copyfileobj(document.file, f)
        new_project.document_path = doc_path

    db.commit()
    db.refresh(new_project)
    return JSONResponse(content={"message": "Project created", "id": new_project.id})

@app.get("/projects/{project_id}/thumbnail")
def get_thumbnail(project_id: int):
    db = SessionLocal()
    project = db.query(Project).filter(Project.id == project_id).first()
    if project and project.thumbnail and os.path.exists(project.thumbnail):
        return FileResponse(project.thumbnail, media_type="image/jpeg")
    return JSONResponse(content={"error": "Thumbnail not found"}, status_code=404)

@app.get("/projects/{project_id}", response_model=ProjectSchema)
def get_project(project_id: int):
    db = SessionLocal()
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project

from docx import Document as DocxDocument

@app.get("/projects/{project_id}/doc_text")
def get_document_text(project_id: int):
    db = SessionLocal()
    project = db.query(Project).filter(Project.id == project_id).first()

    if not project or not project.document_path:
        raise HTTPException(status_code=404, detail="Document not found")

    doc_path = project.document_path
    ext = os.path.splitext(doc_path)[-1].lower()

    try:
        if ext == ".docx":
            doc = DocxDocument(doc_path)
            text = "\n".join([para.text for para in doc.paragraphs])

        elif ext == ".txt" or ext == ".md":
            with open(doc_path, "r", encoding="utf-8") as f:
                text = f.read()

        elif ext == ".pdf":
            with open(doc_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)

        elif ext == ".csv":
            df = pd.read_csv(doc_path)
            text = df.to_string(index=False)

        else:
            raise HTTPException(status_code=415, detail=f"Unsupported document format: {ext}")

        return {"content": text}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read document: {e}")
        
        
@app.delete("/projects/{project_id}")
def delete_project(project_id: int):
    db = SessionLocal()
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Delete files on disk
    project_dir = f"project_files/{project.id}"
    if os.path.exists(project_dir):
        shutil.rmtree(project_dir)

    # Delete from database
    db.delete(project)
    db.commit()

    return {"message": f"Project {project_id} deleted successfully"}


@app.post("/projects/{project_id}/save_doc")
async def save_document(project_id: int, request: Request):
    data = await request.json()
    content = data.get("content")

    db = SessionLocal()
    project = db.query(Project).filter(Project.id == project_id).first()

    if not project or not project.document_path:
        raise HTTPException(status_code=404, detail="Document not found")

    doc_path = project.document_path
    ext = os.path.splitext(doc_path)[-1].lower()

    try:
        if ext == ".docx":
            doc = DocxDocument()
            for line in content.splitlines():
                doc.add_paragraph(line)
            doc.save(doc_path)

        elif ext in [".txt", ".md"]:
            with open(doc_path, "w", encoding="utf-8") as f:
                f.write(content)

        elif ext == ".csv":
            # Save each line as a new row in a one-column CSV
            lines = content.strip().splitlines()
            df = pd.DataFrame({"content": lines})
            df.to_csv(doc_path, index=False)

        else:
            raise HTTPException(status_code=415, detail=f"Unsupported document format: {ext}")

        return {"message": f"Document saved successfully as {ext}"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save document: {e}")

# LLM WORK
# --- LLM and Embedding Configuration ---
# 💬 Option 1: Use local Qwen model and BGE embedding
MODEL_ID = "Qwen/Qwen3-0.6B"
LOCAL_MODEL_DIR = "./models/qwen3-0.6b"
LOCAL_EMBED_DIR = "./models/bge-small-en"

# Load tokenizer for Qwen
tokenizer = AutoTokenizer.from_pretrained(
    MODEL_ID,
    cache_dir=LOCAL_MODEL_DIR
)

# Load Qwen3-4B model with 4-bit quantization
bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.float16,
    bnb_4bit_use_double_quant=False
)

llm_model = AutoModelForCausalLM.from_pretrained(
    MODEL_ID,
    cache_dir=LOCAL_MODEL_DIR,
    device_map="mps",  # or "auto" for CUDA
    torch_dtype=torch.float16,
    trust_remote_code=True
)

# Wrap LLM for LlamaIndex
llm = HuggingFaceLLM(
    model=llm_model,
    tokenizer=tokenizer,
    query_wrapper_prompt="<|im_start|>user\n{query_str}<|im_end|>\n<|im_start|>assistant\n",
    context_window=4096,
    max_new_tokens=1024,
    generate_kwargs={"temperature": 0.1, "do_sample": True}
)

# Embedding model - Choose one
# Option 1: BGE local model
embed_model = HuggingFaceEmbedding(model_name=LOCAL_EMBED_DIR)

# 💬 Option 2: Use OpenAI models
# from openai import OpenAI as OpenAIClient
# import os
# os.environ["OPENAI_API_KEY"] = "sk-..."  # Your OpenAI API Key here
# llm = OpenAI(model="gpt-4o-mini", temperature=0.2)
# embed_model = OpenAIEmbedding(model="text-embedding-3-large")

# Update LlamaIndex settings
Settings.llm = llm
Settings.embed_model = embed_model

# Session memory
chat_sessions = {}

@app.post("/projects/{project_id}/ask")
async def ask_question(project_id: int, request: Request):
    data = await request.json()
    question = data.get("question")

    if not question:
        raise HTTPException(status_code=400, detail="Missing question")

    db = SessionLocal()
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project or not project.document_path or not os.path.exists(project.document_path):
        raise HTTPException(status_code=404, detail="Project document not found")

    docs_dir = os.path.dirname(project.document_path)
    try:
        documents = SimpleDirectoryReader(docs_dir, recursive=True).load_data()
        index = VectorStoreIndex.from_documents(documents)
        query_engine = index.as_query_engine(similarity_top_k=7)

        session_id = str(project_id)
        if session_id not in chat_sessions:
            chat_sessions[session_id] = {"uuid": str(uuid.uuid4()), "context": []}
        session = chat_sessions[session_id]

        session["context"].append({"role": "user", "content": question})

        # --- Use the LlamaIndex query engine to get the response ---
        response = query_engine.query(question)
        response_text = str(response)

        # Qwen-specific post-processing
        if "<think>" in response_text:
            answer_only = re.split(r"</think>\s*", response_text)[-1].strip()
        else:
            answer_only = response_text.strip()
            
        answer = answer_only
        session["context"].append({"role": "assistant", "content": answer})

        return {"question": question, "answer": answer}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi khi xử lý tài liệu: {e}")
